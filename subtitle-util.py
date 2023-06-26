# Simple utility to save as .docx and convert to .srt for external translation

import docx
import sys
from pysubparser import parser

class Translator:
    def __init__(self, input_file_name):
        self.input_file_name = input_file_name
        self.subtitles = parser.parse(input_file_name)
        self.output_file_name = input_file_name.rsplit(".", 1)[0] + ".docx"
        self.doc = docx.Document()
        self.length = 0       
        self.start_list = []
        self.end_list = [] 

    def save_subtitles(self):
        for subtitle in self.subtitles:
            para = self.doc.add_paragraph()
            run = para.add_run(subtitle.text)
            self.length += 1
            self.start_list.append(subtitle.start)
            self.end_list.append(subtitle.end)

    def translate_subtitles(self):
        # wait for translated docx
        print(">> 주의: 파일 번역은 DeepL App이나 웹사이트에서 직접 해주셔야 합니다(로그인 불필요).")
        input(">> " + self.output_file_name + "을 DeepL에서 한국어로 파일 번역한 후 [Enter]를 누르세요...")

        # read translated docx
        # " ko" is appended to the file name by DeepL
        translated_docx = self.output_file_name.rsplit(".", 1)[0] + " ko.docx"
        docx_file = docx.Document(translated_docx)

        # Remove empty line
        translated_texts = []
        for paragraph in docx_file.paragraphs:
            if len(paragraph.text) < 1:
                continue
            translated_texts.append(paragraph.text)

        # check length of translated texts and original texts
        if self.length != len(translated_texts):
            print("오류 자막의 갯수가 맞지 않습니다. SRT의 자막수 vs. 번역된 자막수")
            print(self.length, len(translated_texts))
            return

        self.translated_subtitle = self.output_file_name.rsplit(".", 1)[0] + " ko.srt"
        with open(self.translated_subtitle, "w", encoding="utf-8") as f:
            i = 0
            for start, end, text in zip(self.start_list, self.end_list, translated_texts):
                f.write(f"{i}\n{start} --> {end}\n{text}\n\n")
                i += 1

    def run(self):
        self.save_subtitles()
        self.translate_subtitles()

if __name__ == "__main__":
    if len(sys.argv) > 1:
        input_file_name = sys.argv[1]

        translator = Translator(input_file_name)
        translator.run()
        sys.exit(0)
    else:
        print("사용법: python3 subtitle-util.py <input_file_name>")
        sys.exit(1)
