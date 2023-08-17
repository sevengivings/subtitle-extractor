# Extract subtitle from video, save as .docx for translating and convert to .srt
# https://github.com/sevengivings/subtitle-extractor

# Requires:
# - stable-ts (https://github.com/jianfch/stable-ts) (pip install stable-ts)
# - whisper (https://github.com/openai/whisper) (pip install git+https://github.com/openai/whisper.git )
# - torch + cuda
# - ffmpeg.exe (https://www.ffmpeg.org/) for stable-ts and whisper 
# - python-docx
# - deepl (https://github.com/DeepLcom/deepl-python)

# OS: Windows 10/11 

import os 
import sys
import docx 
import argparse
import torch
import stable_whisper
import whisper 
from whisper.utils import get_writer
import numpy as np
import gettext
import locale 

# stable-ts  
def extract_audio_stable_whisper(model_name, device, use_condition_on_previous_text, use_demucs, use_vad, vad_threshold, is_mel_first, audio_language, input_file_name, output_file_name): 
    # Check if the file exists.
    if not os.path.exists(input_file_name):
        raise FileNotFoundError(f"The file {input_file_name} does not exist.")

    print(f'condition_on_previous_text: {use_condition_on_previous_text}, demucs: {use_demucs}, vad: {use_vad}, vad_threshold: {vad_threshold}, mel_first: {is_mel_first}')

    # Extract the audio from the video.
    model = stable_whisper.load_model(model_name, device=device)
    result = model.transcribe(verbose=True, word_timestamps=False, condition_on_previous_text=use_condition_on_previous_text, \
                                demucs=use_demucs, vad=use_vad, vad_threshold=vad_threshold, mel_first=is_mel_first, \
                                language=audio_language if audio_language != "None" else None, audio=input_file_name)
    result.to_srt_vtt(output_file_name + ".srt", word_level=False)

# Whisper 
def extract_audio_whisper(model_name, device, use_condition_on_previous_text, audio_language, input_file_name):
    # Check if the file exists.
    if not os.path.exists(input_file_name):
        raise FileNotFoundError(f"The file {input_file_name} does not exist.")
    print(f'condition_on_previous_text: {use_condition_on_previous_text}')
    model = whisper.load_model(model_name).to(device)
    temperature = tuple(np.arange(0, 1.0 + 1e-6, 0.2))  # copied from Whisper original code 
    result = model.transcribe(input_file_name, temperature=temperature, verbose=True, word_timestamps=False, condition_on_previous_text=use_condition_on_previous_text, language=audio_language)
    output_dir = os.path.dirname(input_file_name)
    writer = get_writer("srt", output_dir)
    writer(result, input_file_name) 

# srt -> .time, .txt, .docx 
# skip_textlength : 0 = all, 1 = igonore 1 character, 2 = ignore 2 characters
def srt_split(input_file_name, skip_textlength):
    # Check if the file exists.
    if not os.path.exists(input_file_name):
        print(_("Error: No input file: ", input_file_name))
        sys.exit(1)

    # Open the input file.
    with open(input_file_name, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # Create a key-value dictionary to store the subtitle text.
    # key = time sync data, value = line of text
    subtitle_text = {}
    # initialize time_sync_data 
    time_sync_data = ""
    
    # number of line deleted 
    deleted_line = 0  
    # non-duplicated deleted subtitle text list, use set() to remove duplicated item  
    deleted_subtitle_text =  set()

    # ignored number of line 
    ignored_line = 0
    # for ignoring repeated subtitle 
    ignored_subtitle = set()
    # last subtitle text 
    last_subtitle_text = ""
    
    # Iterate over the lines in the file.
    for line in lines: 
        # Ignore empty lines and lines that contain only numbers.
        if line.strip() == "" or line.strip().isdigit():
            time_sync_data = ""
            continue

        # If the line contains "-->", it is time sync data.
        elif line.find("-->") != -1:
            # save time sync data to a varible for later use. 
            time_sync_data = line.strip()

        # Otherwise, the line is subtitle text.
        else:
            # ignore short text under n characters
            if len(line.strip()) > skip_textlength and time_sync_data.find('-->') != -1:
                # Add the line of text to the dictionary.
                value = subtitle_text.get(time_sync_data, "")
                if value is None or value == "":
                    
                    # Ignore repeated subtitles that say the same thing 
                    if last_subtitle_text ==  line.strip():
                        ignored_subtitle.add(line.strip())
                        ignored_line = ignored_line + 1
                        continue
                
                    subtitle_text[time_sync_data] = line.strip()
                    last_subtitle_text = line.strip()
                else: 
                    # If the subtitle text is repeated, combine them into one line.
                    subtitle_text[time_sync_data] = subtitle_text[time_sync_data] + ", " + line.strip() 
                # print(time_sync_data, line)
            else: 
                deleted_line = deleted_line + 1
                deleted_subtitle_text.add(line.strip())

    text_total_length = 0
    
    # Save the subtitle text to a new file.
    # create a new file with the same name as the input file, consider file name contains multiple '.' 
    output_file_name = input_file_name.rsplit(".", 1)[0]
    
    # add extension to the output file name
    f1 = open(output_file_name + ".txt", "w", encoding="utf-8")
    f2 = open(output_file_name + ".time", "w", encoding="utf-8")
    
    # docx file name 
    docx_file_name = output_file_name + ".docx"
    doc = docx.Document()
    doc.save(docx_file_name)
    
    # iterate all subtitle_text and time_sync_data
    for time_sync in subtitle_text:
        f2.write(time_sync + "\n")
        f1.write(subtitle_text[time_sync] + "\n")
        
        para = doc.add_paragraph()
        run = para.add_run(subtitle_text[time_sync])
        
        # CRLF included 
        text_total_length = text_total_length +  len(subtitle_text[time_sync]) + 2 
            
    doc.save(docx_file_name)
    
    # Print the total_length
    print(_("Info: The total length of subtitles is: "), text_total_length)
    
    if len(deleted_subtitle_text) > 0: 
        print(_("Info: The short subtitles are ignored: "), deleted_line)
        output_file_name = input_file_name.rsplit(".", 1)[0] 
        with open(output_file_name + "_short_subtitle_ignored.txt", "w", encoding="utf-8") as f3:
            for text in deleted_subtitle_text:
                f3.write(text + "\n")
                #print(text)

    if len(ignored_subtitle) > 0:
        print(_("Info: The repeated subtitles are ignored: "), ignored_line)
        output_file_name = input_file_name.rsplit(".", 1)[0]
        with open(output_file_name + "_repeated_subtitles_ignored.txt", "w", encoding="utf-8") as f4:
            for text in ignored_subtitle:
                f4.write(text + "\n")
                #print(text)

def docx_to_txt(input_file_name): 
    # Check if the file exists.
    if not os.path.exists(input_file_name):
        raise FileNotFoundError(f"The file {input_file_name} does not exist.")

    # open a docx file and read it
    docx_file = docx.Document(input_file_name)

    # create a new file with the same name as the input file, consider file name contains multiple '.' 
    output_file_name = input_file_name.rsplit(".", 1)[0]

    # save the content of docx file to a new file with the same name as the input file
    with open(output_file_name + ".txt", "w", encoding="utf-8") as f:
        for paragraph in docx_file.paragraphs:
            if len(paragraph.text) < 1:
                print (paragraph.text)
                continue
            f.write(paragraph.text + "\n")
            
    print (output_file_name + _(".txt file is saved"))

def join_srt_files(time_file, text_file, output_file):
    with open(time_file, 'r', encoding='utf-8') as time_input, open(text_file, 'r', encoding='utf-8') as text_input, open(output_file, 'w', encoding='utf-8') as output:
        time_data = time_input.readlines()
        text_data = text_input.readlines()
        if len(time_data) != len(text_data):
            print(_("Error: time and text file must have the same number of lines"))
            return
        for i in range(len(time_data)):
            time_entry = time_data[i].strip()
            text_entry = text_data[i].strip()
            output.write(f"{i}\n")
            output.write(f"{time_entry}\n")
            output.write(f"{text_entry}\n")
            output.write("\n")
    
    print(_("Info: new srt file is saved"), output_file)

def translate_manually(output_file_name, subtitle_language): 
   # Get the translated file name from console if not from DeepL file translation
    print(_("Info: You should translate .docx manullay using DeepL file translation, use "), output_file_name + ".docx")
    # input file name from console  
    file_name = input(_("Input another translated file name or press [Enter] to continue...(") + output_file_name + " " + subtitle_language + _(".docx will be used.): "))
    
    # If input is not given, use default file name.
    if len(file_name) < 1: 
        file_name = output_file_name + " " + subtitle_language + ".docx"

    if not os.path.exists(file_name): 
        print(_("Error: File not found"), file_name)
        sys.exit(1)
    
    # If input file is .txt, docx_to_txt is not used, instead .srt is used to refer to .txt file name.
    if not file_name.endswith(".txt"):
        # Extract .txt from .docx
        docx_to_txt(file_name)
        # "\n>> " + file_name.rsplit(".", 1)[0] + 
        input(_("Press [Enter] to continue... or edit ") + file_name.rsplit(".", 1)[0] + ".txt ")
    
    text_file_name = file_name.rsplit(".", 1)[0]
    
    return text_file_name

def translate_automatically(output_file_name, subtitle_language, deepl_api_key):
    try: 
        import deepl
    except ModuleNotFoundError:
        print(
            "Please install deepl python package by running: pip install --upgrade deepl"
        )
        sys.exit(1)

    translator = deepl.Translator(deepl_api_key)
    input_path = output_file_name + ".docx"
    output_path = output_file_name + " " + subtitle_language + ".docx"
    try: 
        target_lang = subtitle_language.upper()

        # Using translate_document_from_filepath() with file paths 
        translator.translate_document_from_filepath(
            input_path,
            output_path,
            target_lang=target_lang
        )

        # Alternatively you can use translate_document() with file IO objects
        with open(input_path, "rb") as in_file, open(output_path, "wb") as out_file:
            translator.translate_document(
                in_file,
                out_file,
                target_lang=target_lang 
            )

    except deepl.DocumentTranslationException as error:
        # If an error occurs during document translation after the document was
        # already uploaded, a DocumentTranslationException is raised. The
        # document_handle property contains the document handle that may be used to
        # later retrieve the document from the server, or contact DeepL support.
        doc_id = error.document_handle.id
        doc_key = error.document_handle.key
        print(f"Error after uploading ${error}, id: ${doc_id} key: ${doc_key}")
        print(_("DeepL API error: "), error)
        sys.exit(1)    
    except deepl.DeepLException as error:
        # Errors during upload raise a DeepLException
        print(error)    
        print(_("DeepL API error: "), error)
        sys.exit(1)    

    usage = translator.get_usage()
    if usage.any_limit_reached:
        print('Translation limit reached.')
    if usage.character.valid:
        print(
            f"Character usage: {usage.character.count} of {usage.character.limit}")
    if usage.document.valid:
        print(f"Document usage: {usage.document.count} of {usage.document.limit}")

    docx_to_txt(output_path)

    text_file_name = output_path.rsplit(".", 1)[0]

    return text_file_name

if __name__ == "__main__":
    appname = 'subtitle-extractor'
    localedir = './locales'
        
    en_i18n = gettext.translation(appname, localedir, fallback=True, languages=['ko'])  # All messages are in Korean
    en_i18n.install()

    print(_("\nAI subtitle transcription and DeepL translation helper tool\n"))

    parser= argparse.ArgumentParser(formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument("audio", type=str, help=_("audio/video file(s) to transcribe"))
    parser.add_argument("--framework", default="none", help=_("name of the stable-ts or Whisper framework to use"))
    parser.add_argument("--model", default="medium", help=_("tiny, base, small, medium, large model to use"))
    parser.add_argument("--device", default="cuda" if torch.cuda.is_available() else "cpu", help=_("device to use for PyTorch inference"))
    parser.add_argument("--language", type=str, default="None", help=_("language spoken in the audio, specify None to perform language detection"))
    parser.add_argument("--subtitle_language", type=str, default="ko", help=_("subtitle target language need only if you plan to use DeepL file translation"))
    parser.add_argument("--skip_textlength", type=int, default=1, help=_("skip short text in the subtitles, useful for removing meaningless words"))
    parser.add_argument("--condition_on_previous_text", action='store_true',
                        help=_("if True, provide the previous output of the model as a prompt for the next window; "
                             "disabling may make the text inconsistent across windows, "
                             "but the model becomes less prone to getting stuck in a failure loop"))
    # stable_ts only 
    parser.add_argument('--demucs', action='store_true',
                        help=_('stable-ts only, whether to reprocess the audio track with Demucs to isolate vocals/remove noise; '
                             'pip install demucs PySoundFile; '
                             'Demucs official repo: https://github.com/facebookresearch/demucs'))
    parser.add_argument('--vad', action='store_true',
                        help=_('stable-ts only, whether to use Silero VAD to generate timestamp suppression mask; '
                             'pip install silero; '
                             'Official repo: https://github.com/snakers4/silero-vad'))
    parser.add_argument('--vad_threshold', type=float, default=0.2,
                        help=_('stable-ts only, threshold for detecting speech with Silero VAD. (Default: 0.2); '
                             'low threshold reduces false positives for silence detection'))
    parser.add_argument('--mel_first', action='store_true',
                        help=_('stable-ts only, process entire audio track into log-Mel spectrogram first instead in chunks'
                             'if audio is not transcribing properly compared to whisper, at the cost of more memory usage for long audio tracks'))

    args = parser.parse_args().__dict__
    framework: str = args.pop("framework")
    model_name: str = args.pop("model")
    device: str = args.pop("device")
    audio_language: str = args.pop("language")
    subtitle_language: str = args.pop("subtitle_language")
    input_file_name: str = args.pop("audio")
    skip_textlength: int = args.pop("skip_textlength")
    use_condition_on_previous_text: bool = args.pop("condition_on_previous_text")

    # stable_ts only
    use_demucs: bool = args.pop("demucs")
    use_vad: bool = args.pop("vad")
    vad_threshold: float = args.pop("vad_threshold")
    use_mel_first: bool = args.pop("mel_first")    

    print("\nframework: " + framework)
    print("model:" + model_name + "\ndevice:" + device  + "\naudio language:" + audio_language + "\nsubtitle language:" + subtitle_language  + "\nigonore n characters:" + str(skip_textlength) + "\naudio:" + input_file_name)
    print("\nPython version: " + sys.version)
    print("Torch version: " + torch.__version__ + "\n")

    if skip_textlength < 0:
        skip_textlength = 0 

    if not os.path.exists(input_file_name): 
        print(_("The input audio/video file does not exist: "), input_file_name)
        sys.exit(1)

    if audio_language == "None":
        print(_("Info: --audio_language is not specified, performing automatic language detection"))

    # create a new file with the same name as the input file, consider file name contains multiple '.' 
    output_file_name = input_file_name.rsplit(".", 1)[0]

    # slect transcribing method if not input by user 
    if framework == "none":
        # to determine transcribing method, get user input 
        # print ("Input the number of the transcribing method you want to use: 1 for Stable-ts, 2 for Whisper") 
        number_selected = input(_("Input transcribing method. 1 for Stable-ts, 2 for Whisper: "))
        try: 
            number_selected = int(number_selected)
            if number_selected == 1:
                framework = "stable-ts"
            elif number_selected == 2:
                framework = "whisper"
        except ValueError:
            print(_("Error: input numbers only"))
            sys.exit(1)

    # AI speech recognition  
    # Check if the file exists
    if not os.path.exists(output_file_name + ".srt"):           
        if framework == "stable-ts":     
            extract_audio_stable_whisper(model_name, device, use_condition_on_previous_text, use_demucs, use_vad, vad_threshold, use_mel_first, audio_language, input_file_name, output_file_name)
        elif framework == "whisper":
            extract_audio_whisper(model_name, device, use_condition_on_previous_text, audio_language, input_file_name)
    else: 
        print(_("Warning: File already exists"))

    # Separate text and time sync data in .SRT and save subtitle text as .docx for translation. 
    # Removed short sentences and repeated subtitles.  
    srt_split(output_file_name + ".srt", skip_textlength)

    # for DeepL API translation, check if DEEPL_API_KEY is set in environment variables.
    # How to set DeepL API key in PowerShell : Set-Item -Path env:DEEPL_API_KEY -Value "your-id"
    # How to remove DeepL API key in PowerShell : Remove-Item -Path env:DEEPL_API_KEY
    text_file_name = ""
    try: 
        deepl_api_key = os.environ["DEEPL_API_KEY"]
        print(_("Info: DeepL API key is set"))
        print(_("Info: DeepL file translation is started, please wait"))
        text_file_name = translate_automatically(output_file_name, subtitle_language, deepl_api_key)
        print(_("Info: DeepL file translation is done"))        
    except KeyError:
        text_file_name = translate_manually(output_file_name, subtitle_language)
    
    # Change the name of original .srt
    os.rename(output_file_name + ".srt", output_file_name + "_original.srt")

    # Create the final subtitle file joining with .time and .txt.
    join_srt_files(output_file_name + ".time", text_file_name + ".txt", text_file_name + ".srt")
    
    # Change the name of final srt same as video file name 
    print (_("Info: final srt file is saved"))
    if (text_file_name != output_file_name):
        os.rename(text_file_name + ".srt", output_file_name + ".srt")
    
    # Delete intermediate files.
    os.unlink (output_file_name + ".time")
    os.unlink (output_file_name + ".txt")

    # If translated .txt is not provided, there is no ~ ko.docx or ~ ko.txt file.
    # If translated .txt is provided by the user, it is not deleted.
    try: 
        os.unlink (output_file_name + " " + subtitle_language + ".docx")
        os.unlink (output_file_name + " " + subtitle_language + ".txt")
        os.unlink (output_file_name + ".docx")
    except FileNotFoundError: 
        pass

    print(_("Done"))

    sys.exit(0)
